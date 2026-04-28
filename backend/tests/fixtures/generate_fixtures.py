"""Generate test fixture files (PDF, DOCX). Run once to create them."""

import os

FIXTURES_DIR = os.path.dirname(__file__)


def generate_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    path = os.path.join(FIXTURES_DIR, "sample.pdf")
    c = canvas.Canvas(path, pagesize=letter)
    c.drawString(72, 720, "TAIC Test Document")
    c.drawString(72, 700, "This is a test PDF for the RAG pipeline.")
    c.drawString(72, 680, "It contains text that should be extractable by pdfplumber.")
    c.save()
    print(f"Generated: {path}")


def generate_docx():
    from docx import Document

    path = os.path.join(FIXTURES_DIR, "sample.docx")
    doc = Document()
    doc.add_heading("TAIC Test Document", level=1)
    doc.add_paragraph("This is a test DOCX for the RAG pipeline.")
    doc.add_paragraph("It contains paragraphs that should be extractable.")
    doc.save(path)
    print(f"Generated: {path}")


if __name__ == "__main__":
    generate_pdf()
    generate_docx()
