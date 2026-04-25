"""Agent sources endpoints: traceability docs, Notion links, Drive links, sources."""

import io
import json
import logging
import os
import time

from fastapi import APIRouter, Body, Depends, File, HTTPException, UploadFile
from google.cloud import storage
from sqlalchemy.orm import Session

from auth import verify_token
from database import get_db, Agent, AgentShare, Document, DocumentChunk, NotionLink, DriveLink
from helpers.agent_helpers import _user_can_access_agent, _user_can_edit_agent
from helpers.tenant import _get_caller_company_id
from redis_client import get_cached_user
from validation import sanitize_filename

logger = logging.getLogger(__name__)
router = APIRouter()


@router.post("/api/agents/{agent_id}/traceability-docs")
async def upload_traceability_doc(
    agent_id: int, file: UploadFile = File(...), user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Upload a traceability document (no chunking/embedding)."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    allowed_ext = {".pdf", ".txt", ".docx", ".xlsx", ".xls", ".csv"}
    file_ext = "." + file.filename.split(".")[-1].lower() if "." in file.filename else ""
    if file_ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"Unsupported format. Use: {', '.join(allowed_ext)}")

    content = await file.read()
    if len(content) > 20 * 1024 * 1024:  # 20 MB limit
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    # Extract text content
    text_content = ""
    try:
        if file_ext == ".pdf":
            import tempfile

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                from file_loader import load_text_from_pdf

                text_content = load_text_from_pdf(tmp_path) or ""
            finally:
                os.unlink(tmp_path)
        elif file_ext == ".docx":
            import docx

            doc_obj = docx.Document(io.BytesIO(content))
            text_content = "\n".join(p.text for p in doc_obj.paragraphs)
        elif file_ext in (".xlsx", ".xls"):
            import pandas as pd

            df = pd.read_excel(io.BytesIO(content))
            text_content = df.to_string(index=False)
        elif file_ext == ".csv":
            import pandas as pd

            text_content = pd.read_csv(io.BytesIO(content)).to_string(index=False)
        else:
            text_content = content.decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning(f"Could not extract text from traceability doc: {e}")
        text_content = ""

    # Upload to GCS
    gcs_url = None
    try:
        GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME", "applydi-documents")
        gcs_client = storage.Client()
        bucket = gcs_client.bucket(GCS_BUCKET_NAME)
        gcs_filename = f"traceability/{int(time.time())}_{sanitize_filename(file.filename)}"
        blob = bucket.blob(gcs_filename)
        blob.upload_from_string(content)
        gcs_url = blob.public_url
    except Exception as e:
        logger.warning(f"GCS upload failed for traceability doc: {e}")

    doc = Document(
        filename=file.filename,
        content=text_content,
        user_id=int(user_id),
        agent_id=agent_id,
        company_id=_get_caller_company_id(user_id, db),
        gcs_url=gcs_url,
        document_type="traceability",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {"document": {"id": doc.id, "filename": doc.filename, "created_at": doc.created_at.isoformat()}}


@router.get("/api/agents/{agent_id}/traceability-docs")
async def list_traceability_docs(agent_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    """List traceability documents for an agent."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    docs = (
        db.query(Document)
        .filter(Document.agent_id == agent_id, Document.document_type == "traceability")
        .order_by(Document.created_at.desc())
        .all()
    )

    return {"documents": [{"id": d.id, "filename": d.filename, "created_at": d.created_at.isoformat()} for d in docs]}


@router.delete("/api/agents/{agent_id}/traceability-docs/{doc_id}")
async def delete_traceability_doc(
    agent_id: int, doc_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Delete a traceability document."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    doc = (
        db.query(Document)
        .filter(Document.id == doc_id, Document.agent_id == agent_id, Document.document_type == "traceability")
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete from GCS if applicable
    if doc.gcs_url:
        try:
            GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME", "applydi-documents")
            gcs_client = storage.Client()
            bucket = gcs_client.bucket(GCS_BUCKET_NAME)
            blob_name = doc.gcs_url.split(f"{GCS_BUCKET_NAME}/")[-1] if GCS_BUCKET_NAME in (doc.gcs_url or "") else None
            if blob_name:
                bucket.blob(blob_name).delete()
        except Exception as e:
            logger.warning(f"Could not delete GCS blob for traceability doc {doc_id}: {e}")

    db.delete(doc)
    db.commit()

    return {"message": "Document deleted"}


## ---- Notion Links Endpoints ----


@router.post("/api/agents/{agent_id}/notion-links")
async def add_notion_link(
    agent_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db), body: dict = Body(...)
):
    """Link a Notion page or database to an agent."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    url = body.get("url", "").strip()
    resource_type = body.get("type", "page").strip()
    if resource_type not in ("page", "database"):
        raise HTTPException(status_code=400, detail="type must be 'page' or 'database'")
    if not url:
        raise HTTPException(status_code=400, detail="url is required")

    from notion_client import extract_notion_id, fetch_page_title, fetch_database_title, get_notion_token

    user = get_cached_user(user_id, db)
    user_company_id = user.company_id if user else None
    if not get_notion_token(user_company_id):
        raise HTTPException(
            status_code=503, detail="Notion integration is not configured. Ask your organization owner to configure it."
        )

    try:
        notion_id = extract_notion_id(url)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    try:
        if resource_type == "page":
            label = fetch_page_title(notion_id, company_id=user_company_id)
        else:
            label = fetch_database_title(notion_id, company_id=user_company_id)
    except Exception as e:
        logger.warning(f"Could not fetch Notion title for {notion_id}: {e}")
        raise HTTPException(status_code=400, detail=f"Could not access Notion resource: {e}")

    link = NotionLink(
        agent_id=agent_id,
        company_id=_get_caller_company_id(user_id, db),
        notion_resource_id=notion_id,
        resource_type=resource_type,
        label=label,
    )
    db.add(link)
    db.commit()
    db.refresh(link)

    return {
        "link": {
            "id": link.id,
            "notion_resource_id": link.notion_resource_id,
            "resource_type": link.resource_type,
            "label": link.label,
            "created_at": link.created_at.isoformat(),
        }
    }


@router.get("/api/agents/{agent_id}/notion-links")
async def list_notion_links(agent_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    """List Notion links for an agent."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    links = db.query(NotionLink).filter(NotionLink.agent_id == agent_id).order_by(NotionLink.created_at.desc()).all()
    # Check which links have been ingested into RAG
    ingested_ids = set()
    if links:
        link_ids = [l.id for l in links]
        ingested_docs = db.query(Document.notion_link_id).filter(Document.notion_link_id.in_(link_ids)).all()
        ingested_ids = {d[0] for d in ingested_docs}
    return {
        "links": [
            {
                "id": l.id,
                "notion_resource_id": l.notion_resource_id,
                "resource_type": l.resource_type,
                "label": l.label,
                "created_at": l.created_at.isoformat(),
                "ingested": l.id in ingested_ids,
            }
            for l in links
        ]
    }


@router.delete("/api/agents/{agent_id}/notion-links/{link_id}")
async def delete_notion_link(
    agent_id: int, link_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Delete a Notion link from an agent."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    link = db.query(NotionLink).filter(NotionLink.id == link_id, NotionLink.agent_id == agent_id).first()
    if not link:
        raise HTTPException(status_code=404, detail="Notion link not found")

    # Also delete any RAG document ingested from this Notion link
    ingested_doc = db.query(Document).filter(Document.notion_link_id == link_id).first()
    if ingested_doc:
        db.delete(ingested_doc)

    db.delete(link)
    db.commit()
    return {"message": "Notion link deleted"}


@router.post("/api/agents/{agent_id}/notion-links/{link_id}/preview")
async def preview_notion_link(
    agent_id: int, link_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Preview the live content of a Notion link."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    link = db.query(NotionLink).filter(NotionLink.id == link_id, NotionLink.agent_id == agent_id).first()
    if not link:
        raise HTTPException(status_code=404, detail="Notion link not found")

    from notion_client import fetch_page_content, fetch_database_entries, blocks_to_text, database_entries_to_text

    user = get_cached_user(user_id, db)
    user_company_id = user.company_id if user else None

    try:
        if link.resource_type == "page":
            blocks = fetch_page_content(link.notion_resource_id, company_id=user_company_id)
            content = blocks_to_text(blocks)
        else:
            entries = fetch_database_entries(link.notion_resource_id, company_id=user_company_id)
            content = database_entries_to_text(entries)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch Notion content: {e}")

    return {"label": link.label, "resource_type": link.resource_type, "content": content}


@router.get("/api/agents/{agent_id}/sources")
async def get_agent_sources(agent_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    """Consolidated sources page: documents + notion links + permissions."""
    agent = _user_can_access_agent(int(user_id), agent_id, db)

    # Check edit permission
    can_edit = agent.user_id == int(user_id)
    if not can_edit:
        share = (
            db.query(AgentShare)
            .filter(AgentShare.agent_id == agent_id, AgentShare.user_id == int(user_id), AgentShare.can_edit == True)
            .first()
        )
        can_edit = share is not None

    # RAG documents for this agent
    docs = (
        db.query(Document)
        .filter(Document.agent_id == agent_id, Document.document_type == "rag")
        .order_by(Document.created_at.desc())
        .all()
    )

    # Notion links
    notion_links = (
        db.query(NotionLink).filter(NotionLink.agent_id == agent_id).order_by(NotionLink.created_at.desc()).all()
    )

    # Set of notion_link_ids that already have a Document
    ingested_link_ids = set()
    for d in docs:
        if d.notion_link_id:
            ingested_link_ids.add(d.notion_link_id)

    # Drive links
    drive_links = db.query(DriveLink).filter(DriveLink.agent_id == agent_id).order_by(DriveLink.created_at.desc()).all()

    return {
        "agent_name": agent.name,
        "documents": [
            {
                "id": d.id,
                "filename": d.filename,
                "created_at": d.created_at.isoformat() if d.created_at else None,
                "has_file": bool(d.gcs_url),
                "notion_link_id": d.notion_link_id,
                "drive_link_id": getattr(d, "drive_link_id", None),
                "source_url": getattr(d, "source_url", None),
            }
            for d in docs
        ],
        "notion_links": [
            {
                "id": nl.id,
                "label": nl.label,
                "resource_type": nl.resource_type,
                "created_at": nl.created_at.isoformat() if nl.created_at else None,
                "ingested": nl.id in ingested_link_ids,
            }
            for nl in notion_links
        ],
        "drive_links": [
            {
                "id": dl.id,
                "drive_folder_id": dl.drive_folder_id,
                "label": dl.label,
                "created_at": dl.created_at.isoformat() if dl.created_at else None,
                "ingested_count": db.query(Document).filter(Document.drive_link_id == dl.id).count(),
            }
            for dl in drive_links
        ],
        "can_edit": can_edit,
    }


@router.post("/api/agents/{agent_id}/notion-links/{link_id}/ingest")
async def ingest_notion_link(
    agent_id: int, link_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Ingest a Notion link's content into the RAG pipeline."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    link = db.query(NotionLink).filter(NotionLink.id == link_id, NotionLink.agent_id == agent_id).first()
    if not link:
        raise HTTPException(status_code=404, detail="Notion link not found")

    # Check not already ingested
    existing = db.query(Document).filter(Document.notion_link_id == link_id).first()
    if existing:
        raise HTTPException(status_code=409, detail="This Notion link has already been ingested")

    from notion_client import fetch_page_content, fetch_database_entries, blocks_to_text, database_entries_to_text

    user = get_cached_user(user_id, db)
    user_company_id = user.company_id if user else None

    try:
        if link.resource_type == "page":
            blocks = fetch_page_content(link.notion_resource_id, company_id=user_company_id)
            text_content = blocks_to_text(blocks)
        else:
            entries = fetch_database_entries(link.notion_resource_id, company_id=user_company_id)
            text_content = database_entries_to_text(entries)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch Notion content: {e}")

    if not text_content or not text_content.strip():
        raise HTTPException(status_code=400, detail="Notion resource returned empty content")

    from rag_engine import ingest_text_content

    safe_label = (link.label or "untitled").replace("/", "_")[:80]
    filename = f"notion_{link.resource_type}_{safe_label}.txt"

    # Upload text to GCS so the document is downloadable
    gcs_url = None
    try:
        from google.cloud import storage as gcs_storage
        import time as _time

        bucket_name = os.getenv("GCS_BUCKET_NAME", "applydi-documents")
        gcs_client = gcs_storage.Client()
        bucket = gcs_client.bucket(bucket_name)
        gcs_filename = f"{int(_time.time())}_{filename.replace(' ', '_')}"
        blob = bucket.blob(gcs_filename)
        blob.upload_from_string(text_content.encode("utf-8"), content_type="text/plain; charset=utf-8")
        gcs_url = blob.public_url
    except Exception as e:
        logger.warning(f"Could not upload Notion text to GCS: {e}")

    doc_id = ingest_text_content(
        text_content=text_content,
        filename=filename,
        user_id=int(user_id),
        agent_id=agent_id,
        db=db,
        gcs_url=gcs_url,
        notion_link_id=link.id,
        company_id=user_company_id,
    )

    doc = db.query(Document).filter(Document.id == doc_id).first()
    chunk_count = len(doc.chunks) if doc else 0

    return {
        "document_id": doc_id,
        "filename": filename,
        "chunk_count": chunk_count,
        "message": "Notion content ingested successfully",
    }


@router.post("/api/agents/{agent_id}/notion-links/{link_id}/resync")
async def resync_notion_link(
    agent_id: int, link_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Re-sync a Notion link: delete existing document then re-ingest from Notion."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    link = db.query(NotionLink).filter(NotionLink.id == link_id, NotionLink.agent_id == agent_id).first()
    if not link:
        raise HTTPException(status_code=404, detail="Notion link not found")

    # Delete existing ingested document if any
    existing = db.query(Document).filter(Document.notion_link_id == link_id).first()
    if existing:
        db.delete(existing)
        db.commit()

    # Re-fetch content from Notion
    from notion_client import fetch_page_content, fetch_database_entries, blocks_to_text, database_entries_to_text

    user = get_cached_user(user_id, db)
    user_company_id = user.company_id if user else None

    try:
        if link.resource_type == "page":
            blocks = fetch_page_content(link.notion_resource_id, company_id=user_company_id)
            text_content = blocks_to_text(blocks)
        else:
            entries = fetch_database_entries(link.notion_resource_id, company_id=user_company_id)
            text_content = database_entries_to_text(entries)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch Notion content: {e}")

    if not text_content or not text_content.strip():
        raise HTTPException(status_code=400, detail="Notion resource returned empty content")

    from rag_engine import ingest_text_content

    safe_label = (link.label or "untitled").replace("/", "_")[:80]
    filename = f"notion_{link.resource_type}_{safe_label}.txt"

    gcs_url = None
    try:
        from google.cloud import storage as gcs_storage
        import time as _time

        bucket_name = os.getenv("GCS_BUCKET_NAME", "applydi-documents")
        gcs_client = gcs_storage.Client()
        bucket = gcs_client.bucket(bucket_name)
        gcs_filename = f"{int(_time.time())}_{filename.replace(' ', '_')}"
        blob = bucket.blob(gcs_filename)
        blob.upload_from_string(text_content.encode("utf-8"), content_type="text/plain; charset=utf-8")
        gcs_url = blob.public_url
    except Exception as e:
        logger.warning(f"Could not upload Notion text to GCS: {e}")

    doc_id = ingest_text_content(
        text_content=text_content,
        filename=filename,
        user_id=int(user_id),
        agent_id=agent_id,
        db=db,
        gcs_url=gcs_url,
        notion_link_id=link.id,
        company_id=user_company_id,
    )

    doc = db.query(Document).filter(Document.id == doc_id).first()
    chunk_count = len(doc.chunks) if doc else 0

    logger.info(f"Notion link {link_id} re-synced for agent {agent_id} by user {user_id}")
    return {
        "document_id": doc_id,
        "filename": filename,
        "chunk_count": chunk_count,
        "message": "Notion content re-synced successfully",
    }


# ── Google Drive integration ──────────────────────────────────────────────


@router.post("/api/agents/{agent_id}/drive-links")
async def add_drive_link(
    agent_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db), body: dict = Body(...)
):
    """Link a Google Drive folder to an agent."""
    agent = _user_can_edit_agent(int(user_id), agent_id, db)

    url = body.get("url", "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="url is required")

    from google_drive_client import extract_drive_folder_id, fetch_folder_name

    try:
        folder_id = extract_drive_folder_id(url)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    try:
        label = fetch_folder_name(folder_id, agent_id=agent_id, db=db)
    except Exception as e:
        logger.warning(f"Could not access Drive folder {folder_id}: {e}")
        raise HTTPException(
            status_code=400,
            detail=f"Cannot access this Drive folder. Make sure it is shared with the backend service account. ({e})",
        )

    link = DriveLink(
        agent_id=agent_id,
        company_id=_get_caller_company_id(user_id, db),
        drive_folder_id=folder_id,
        label=label,
    )
    db.add(link)
    db.commit()
    db.refresh(link)

    return {
        "link": {
            "id": link.id,
            "drive_folder_id": link.drive_folder_id,
            "label": link.label,
            "created_at": link.created_at.isoformat(),
        }
    }


@router.get("/api/agents/{agent_id}/drive-links")
async def list_drive_links(agent_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    """List Google Drive links for an agent."""
    _user_can_edit_agent(int(user_id), agent_id, db)

    links = db.query(DriveLink).filter(DriveLink.agent_id == agent_id).order_by(DriveLink.created_at.desc()).all()

    # Count ingested documents per link
    result = []
    for link in links:
        ingested_count = db.query(Document).filter(Document.drive_link_id == link.id).count()
        result.append(
            {
                "id": link.id,
                "drive_folder_id": link.drive_folder_id,
                "label": link.label,
                "created_at": link.created_at.isoformat(),
                "ingested_count": ingested_count,
            }
        )

    return {"links": result}


@router.delete("/api/agents/{agent_id}/drive-links/{link_id}")
async def delete_drive_link(
    agent_id: int, link_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Delete a Drive link and all its ingested documents."""
    _user_can_edit_agent(int(user_id), agent_id, db)

    link = db.query(DriveLink).filter(DriveLink.id == link_id, DriveLink.agent_id == agent_id).first()
    if not link:
        raise HTTPException(status_code=404, detail="Drive link not found")

    # Delete all documents ingested from this link (cascades to chunks)
    docs = db.query(Document).filter(Document.drive_link_id == link_id).all()
    for doc in docs:
        db.delete(doc)

    db.delete(link)
    db.commit()
    return {"message": "Drive link deleted"}


@router.post("/api/agents/{agent_id}/drive-links/{link_id}/ingest")
async def ingest_drive_link(
    agent_id: int, link_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Ingest all files from a Drive folder into the RAG pipeline."""
    _user_can_edit_agent(int(user_id), agent_id, db)

    link = db.query(DriveLink).filter(DriveLink.id == link_id, DriveLink.agent_id == agent_id).first()
    if not link:
        raise HTTPException(status_code=404, detail="Drive link not found")

    from google_drive_client import get_drive_service, list_folder_files, extract_text_from_drive_file, SUPPORTED_MIMES
    from rag_engine import ingest_text_content

    try:
        service = get_drive_service(agent_id=agent_id, db=db)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))

    user = get_cached_user(user_id, db)
    user_company_id = user.company_id if user else None

    try:
        files = list_folder_files(link.drive_folder_id, agent_id=agent_id, db=db)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to list Drive folder: {e}")

    files_processed = 0
    files_skipped = 0
    total_chunks = 0

    for f in files:
        if f["mimeType"] not in SUPPORTED_MIMES:
            files_skipped += 1
            continue

        text = extract_text_from_drive_file(service, f["id"], f["name"], f["mimeType"])
        if not text or not text.strip():
            files_skipped += 1
            continue

        try:
            safe_name = f["name"].replace("/", "_")[:100]
            filename = f"drive_{safe_name}"

            doc_id = ingest_text_content(
                text_content=text,
                filename=filename,
                user_id=int(user_id),
                agent_id=agent_id,
                db=db,
                drive_link_id=link.id,
                drive_file_id=f["id"],
                company_id=user_company_id,
            )
            doc = db.query(Document).filter(Document.id == doc_id).first()
            total_chunks += len(doc.chunks) if doc else 0
            files_processed += 1
        except Exception as e:
            logger.warning(f"Failed to ingest Drive file {f['name']}: {e}")
            files_skipped += 1

    return {"files_processed": files_processed, "files_skipped": files_skipped, "chunk_count": total_chunks}


@router.post("/api/agents/{agent_id}/drive-links/{link_id}/resync")
async def resync_drive_link(
    agent_id: int, link_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    """Sync new files from a Drive folder (skip already ingested files)."""
    _user_can_edit_agent(int(user_id), agent_id, db)

    link = db.query(DriveLink).filter(DriveLink.id == link_id, DriveLink.agent_id == agent_id).first()
    if not link:
        raise HTTPException(status_code=404, detail="Drive link not found")

    from google_drive_client import get_drive_service, list_folder_files, extract_text_from_drive_file, SUPPORTED_MIMES
    from rag_engine import ingest_text_content

    try:
        service = get_drive_service(agent_id=agent_id, db=db)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))

    user = get_cached_user(user_id, db)
    user_company_id = user.company_id if user else None

    try:
        files = list_folder_files(link.drive_folder_id, agent_id=agent_id, db=db)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to list Drive folder: {e}")

    # Get already ingested file IDs
    existing_file_ids = set(
        row[0]
        for row in db.query(Document.drive_file_id)
        .filter(Document.drive_link_id == link_id, Document.drive_file_id.isnot(None))
        .all()
    )

    new_files_added = 0
    total_chunks = 0

    for f in files:
        if f["id"] in existing_file_ids:
            continue
        if f["mimeType"] not in SUPPORTED_MIMES:
            continue

        text = extract_text_from_drive_file(service, f["id"], f["name"], f["mimeType"])
        if not text or not text.strip():
            continue

        try:
            safe_name = f["name"].replace("/", "_")[:100]
            filename = f"drive_{safe_name}"

            doc_id = ingest_text_content(
                text_content=text,
                filename=filename,
                user_id=int(user_id),
                agent_id=agent_id,
                db=db,
                drive_link_id=link.id,
                drive_file_id=f["id"],
                company_id=user_company_id,
            )
            doc = db.query(Document).filter(Document.id == doc_id).first()
            total_chunks += len(doc.chunks) if doc else 0
            new_files_added += 1
        except Exception as e:
            logger.warning(f"Failed to ingest Drive file {f['name']} during resync: {e}")

    return {"new_files_added": new_files_added, "chunk_count": total_chunks}
