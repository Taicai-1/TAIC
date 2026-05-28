"""Recap CRUD and action endpoints."""

import io
import json
import logging
import os
import time
from datetime import datetime

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy.orm import Session

from auth import verify_token
from database import get_db, Agent, Document, Recap, RecapDocument, WeeklyRecapLog, User
from helpers.tenant import _get_caller_company_id
from validation import sanitize_filename

logger = logging.getLogger(__name__)
router = APIRouter()


class RecapCreate(BaseModel):
    name: str
    enabled: bool = True
    frequency: str = "weekly"
    hour: int = 9
    prompt: str | None = None
    recipients: list[str] | None = None


class RecapUpdate(BaseModel):
    name: str | None = None
    enabled: bool | None = None
    frequency: str | None = None
    hour: int | None = None
    prompt: str | None = None
    recipients: list[str] | None = None


def _get_agent_for_user(agent_id: int, user_id: int, db: Session) -> Agent:
    agent = db.query(Agent).filter(Agent.id == agent_id, Agent.user_id == user_id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")
    return agent


def _get_recap_for_user(recap_id: int, user_id: int, db: Session) -> Recap:
    recap = db.query(Recap).filter(Recap.id == recap_id).first()
    if not recap:
        raise HTTPException(status_code=404, detail="Recap not found")
    agent = db.query(Agent).filter(Agent.id == recap.agent_id, Agent.user_id == user_id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Recap not found")
    return recap


def _recap_to_dict(recap: Recap) -> dict:
    recipients = []
    if recap.recipients:
        try:
            recipients = json.loads(recap.recipients)
        except (json.JSONDecodeError, TypeError):
            pass
    return {
        "id": recap.id,
        "agent_id": recap.agent_id,
        "name": recap.name,
        "enabled": recap.enabled,
        "frequency": recap.frequency,
        "hour": recap.hour,
        "prompt": recap.prompt,
        "recipients": recipients,
        "created_at": recap.created_at.isoformat() if recap.created_at else None,
        "updated_at": recap.updated_at.isoformat() if recap.updated_at else None,
        "document_count": len(recap.recap_documents),
    }


@router.get("/api/agents/{agent_id}/recaps")
async def list_recaps(agent_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    agent = _get_agent_for_user(agent_id, int(user_id), db)
    recaps = db.query(Recap).filter(Recap.agent_id == agent.id).order_by(Recap.created_at.asc()).all()
    return {"recaps": [_recap_to_dict(r) for r in recaps]}


@router.post("/api/agents/{agent_id}/recaps")
async def create_recap(
    agent_id: int, body: RecapCreate, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    agent = _get_agent_for_user(agent_id, int(user_id), db)
    caller_company_id = _get_caller_company_id(user_id, db)

    if body.frequency not in ("daily", "weekly", "monthly"):
        raise HTTPException(status_code=400, detail="frequency must be daily, weekly, or monthly")
    if not (0 <= body.hour <= 23):
        raise HTTPException(status_code=400, detail="hour must be between 0 and 23")

    recap = Recap(
        agent_id=agent.id,
        company_id=caller_company_id,
        name=body.name,
        enabled=body.enabled,
        frequency=body.frequency,
        hour=body.hour,
        prompt=body.prompt if body.prompt and body.prompt.strip() else None,
        recipients=json.dumps(body.recipients) if body.recipients else None,
    )
    db.add(recap)
    db.commit()
    db.refresh(recap)

    return {"recap": _recap_to_dict(recap)}


@router.put("/api/agents/{agent_id}/recaps/{recap_id}")
async def update_recap(
    agent_id: int, recap_id: int, body: RecapUpdate, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    _get_agent_for_user(agent_id, int(user_id), db)
    recap = _get_recap_for_user(recap_id, int(user_id), db)

    if body.name is not None:
        recap.name = body.name
    if body.enabled is not None:
        recap.enabled = body.enabled
    if body.frequency is not None:
        if body.frequency not in ("daily", "weekly", "monthly"):
            raise HTTPException(status_code=400, detail="frequency must be daily, weekly, or monthly")
        recap.frequency = body.frequency
    if body.hour is not None:
        if not (0 <= body.hour <= 23):
            raise HTTPException(status_code=400, detail="hour must be between 0 and 23")
        recap.hour = body.hour
    if body.prompt is not None:
        recap.prompt = body.prompt if body.prompt.strip() else None
    if body.recipients is not None:
        recap.recipients = json.dumps(body.recipients) if body.recipients else None

    recap.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(recap)

    return {"recap": _recap_to_dict(recap)}


@router.delete("/api/agents/{agent_id}/recaps/{recap_id}")
async def delete_recap(
    agent_id: int, recap_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)
):
    _get_agent_for_user(agent_id, int(user_id), db)
    recap = _get_recap_for_user(recap_id, int(user_id), db)
    db.delete(recap)
    db.commit()
    return {"message": "Recap deleted"}


@router.get("/api/recaps/{recap_id}/documents")
async def list_recap_documents(recap_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    recap = _get_recap_for_user(recap_id, int(user_id), db)
    docs = (
        db.query(Document)
        .join(RecapDocument, RecapDocument.document_id == Document.id)
        .filter(RecapDocument.recap_id == recap.id)
        .order_by(Document.created_at.desc())
        .all()
    )

    return {
        "documents": [
            {
                "document_id": doc.id,
                "filename": doc.filename,
                "created_at": doc.created_at.isoformat() if doc.created_at else None,
            }
            for doc in docs
        ]
    }


@router.delete("/api/recaps/{recap_id}/documents/{document_id}")
async def remove_recap_document(
    recap_id: int,
    document_id: int,
    user_id: str = Depends(verify_token),
    db: Session = Depends(get_db),
):
    recap = _get_recap_for_user(recap_id, int(user_id), db)
    rd = (
        db.query(RecapDocument)
        .filter(RecapDocument.recap_id == recap.id, RecapDocument.document_id == document_id)
        .first()
    )
    if not rd:
        raise HTTPException(status_code=404, detail="Document not linked to this recap")
    db.delete(rd)
    db.commit()
    return {"message": "Document removed from recap"}


@router.post("/api/recaps/{recap_id}/documents/upload")
async def upload_recap_document(
    recap_id: int,
    file: UploadFile = File(...),
    user_id: str = Depends(verify_token),
    db: Session = Depends(get_db),
):
    """Upload a traceability document and link it to a specific recap."""
    recap = _get_recap_for_user(recap_id, int(user_id), db)
    agent = db.query(Agent).filter(Agent.id == recap.agent_id).first()
    caller_company_id = _get_caller_company_id(user_id, db)

    allowed_ext = {".pdf", ".txt", ".docx", ".xlsx", ".xls", ".csv", ".json"}
    file_ext = "." + file.filename.split(".")[-1].lower() if "." in file.filename else ""
    if file_ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"Unsupported format. Use: {', '.join(allowed_ext)}")

    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    # Extract text content
    text_content = ""
    try:
        if file_ext == ".pdf":
            import tempfile

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                from file_loader import load_text_from_pdf

                text_content = load_text_from_pdf(tmp_path) or ""
            finally:
                os.unlink(tmp_path)
        elif file_ext == ".docx":
            import docx

            doc_obj = docx.Document(io.BytesIO(content))
            text_content = "\n".join(p.text for p in doc_obj.paragraphs)
        elif file_ext in (".xlsx", ".xls"):
            import pandas as pd

            df = pd.read_excel(io.BytesIO(content))
            text_content = df.to_string(index=False)
        elif file_ext == ".csv":
            import pandas as pd

            text_content = pd.read_csv(io.BytesIO(content)).to_string(index=False)
        else:
            text_content = content.decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning(f"Could not extract text from traceability doc: {e}")
        text_content = ""

    # Upload to GCS
    gcs_url = None
    try:
        from google.cloud import storage

        GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME", "applydi-documents")
        gcs_client = storage.Client()
        bucket = gcs_client.bucket(GCS_BUCKET_NAME)
        gcs_filename = f"traceability/{int(time.time())}_{sanitize_filename(file.filename)}"
        blob = bucket.blob(gcs_filename)
        blob.upload_from_string(content)
        gcs_url = blob.public_url
    except Exception as e:
        logger.warning(f"GCS upload failed for traceability doc: {e}")

    doc = Document(
        filename=file.filename,
        content=text_content,
        user_id=int(user_id),
        agent_id=agent.id,
        company_id=caller_company_id,
        gcs_url=gcs_url,
        document_type="traceability",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    rd = RecapDocument(
        recap_id=recap.id,
        document_id=doc.id,
        included=True,
        company_id=caller_company_id,
    )
    db.add(rd)
    db.commit()

    return {
        "document": {
            "document_id": doc.id,
            "filename": doc.filename,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
        }
    }


@router.post("/api/recaps/{recap_id}/preview")
async def recap_preview(recap_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    recap = _get_recap_for_user(recap_id, int(user_id), db)
    agent = db.query(Agent).filter(Agent.id == recap.agent_id).first()

    from weekly_recap import (
        fetch_weekly_messages,
        fetch_recap_traceability_documents,
        fetch_notion_content,
        build_recap_prompt,
        get_model_id_for_agent,
        FREQUENCY_DAYS,
    )
    from email_service import generate_recap_html
    from openai_client import get_chat_response as _get_chat_response

    days_back = FREQUENCY_DAYS.get(recap.frequency, 7)
    messages = fetch_weekly_messages(agent.id, db, days_back=days_back)
    docs = fetch_recap_traceability_documents(recap.id, db, days_back=days_back)
    notion_pages = fetch_notion_content(agent.id, db)

    if not messages and not docs and not notion_pages:
        return {"status": "no_data", "message": "No data for this period", "html": None}

    prompt_messages = build_recap_prompt(
        agent, messages, docs, notion_pages, frequency=recap.frequency, custom_prompt=recap.prompt
    )
    model_id = get_model_id_for_agent(agent)
    recap_content = _get_chat_response(prompt_messages, model_id=model_id)
    html = generate_recap_html(agent.name, recap_content, recap_name=recap.name)

    return {
        "status": "success",
        "html": html,
        "message_count": len(messages),
        "doc_count": len(docs),
        "notion_count": len(notion_pages),
    }


@router.post("/api/recaps/{recap_id}/send")
async def recap_send(recap_id: int, user_id: str = Depends(verify_token), db: Session = Depends(get_db)):
    recap = _get_recap_for_user(recap_id, int(user_id), db)
    if not recap.enabled:
        raise HTTPException(status_code=400, detail="Recap is not enabled")

    from weekly_recap import process_recap

    result = process_recap(recap, db)
    return result
